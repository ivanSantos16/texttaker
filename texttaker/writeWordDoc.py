import json
import docx
import os

def writeWordDoc(json_file: str, msdocpath: str):
    '''
    This function takes the json file and writes it to a word document
    
    Parameters:
    json_file (str): The path to json file that have the data to be written to a word document
    
    '''

    data = json.load(open(json_file, encoding="utf8"))
    for pdf_name in data:
        doc = docx.Document()
        title = pdf_name.split('.')[0] + ' | Compute Time of The File: ' + str(data[pdf_name]["compute_time(s)"]) + ' seconds'
        doc.add_heading(title, 0)
        doc.add_paragraph(data[pdf_name]["text"])
        if not os.path.exists(msdocpath):
            os.makedirs(msdocpath)
        doc.save(msdocpath + pdf_name.split('.')[0] + '.docx')


# writeWordDoc(json_file = './data/text_extracted.json')

# Check links:
# https://stackoverflow.com/questions/9505898/conditional-command-line-arguments-in-python-using-argparse
# https://stackoverflow.com/questions/30896982/argparse-optional-value-for-argument       

